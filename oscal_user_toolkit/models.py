"""
models.py
=========
This file contains all the "data" logic for the OSCAL User Toolkit.
It has NO graphical user interface (GUI) code at all — it only deals
with reading, writing, and converting data.

Keeping data logic separate from GUI code is good practice because:
  - It is easier to test (you can run these functions without opening a window)
  - It is easier to read (each file has one clear job)
  - If you ever change the GUI library, this file stays the same

OSCAL is an open standard for describing security controls and plans.
Files are stored as JSON (JavaScript Object Notation), a plain-text
format that looks like Python dictionaries and lists.
"""

# ── Standard library imports ──────────────────────────────────────────────────
# These modules are built into Python — no installation needed.

import json        # Reads and writes JSON files
import re          # Regular expressions — imported here at module level (L2 fix)
import uuid        # Generates universally unique identifiers (UUIDs)
import zipfile
from datetime import datetime, timezone   # Used to get the current date/time
from pathlib import Path                  # Cross-platform file path handling

try:
    # jsonschema validates a JSON document against a schema (a set of rules).
    # It is an optional dependency — if not installed, schema validation is
    # silently skipped. We record availability in a flag so other code can
    # check it without repeating the try/except.
    import jsonschema
    _JSONSCHEMA_AVAILABLE = True
except ImportError:
    _JSONSCHEMA_AVAILABLE = False

try:
    # python-docx builds Microsoft Word (.docx) files programmatically.
    # Also optional — if missing, the Export DOCX button will show an error.
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


# ── OSCAL version constant ────────────────────────────────────────────────────
# The OSCAL version this toolkit targets. Used as a fallback when the
# calling tab does not supply a version getter. Keeping it as a named
# constant (rather than a bare string) means a single edit updates every
# place that uses it. (M1 fix)
DEFAULT_OSCAL_VERSION = "1.2.2"

# ── Pre-compiled UUID regex ───────────────────────────────────────────────────
# Compiled once at module load time rather than on every function call.
# Matches the standard 8-4-4-4-12 hex digit UUID format. (L9 fix)
_UUID_RE = re.compile(
    r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$',
    re.IGNORECASE,
)

# ── Shared OSCAL 1.2.2 enumeration constants ──────────────────────────────────
# These lists define the allowed values for various OSCAL fields.
# They are defined here once and imported by the tab UI files to ensure
# both the AR editor and POA&M editor always show the same options. (L1 fix)

# Methods used to gather assessment evidence (OSCAL observation.methods)
OBSERVATION_METHODS = ["EXAMINE", "INTERVIEW", "TEST", "UNKNOWN"]

# Types of observation (OSCAL observation.types — common values)
OBSERVATION_TYPES = [
    "", "ssp-statement-issue", "control-objective-not-met",
    "finding", "historic", "tool-generated",
]

# Risk lifecycle status values (OSCAL risk.status)
RISK_STATUSES = [
    "open", "investigating", "remediating", "deviation-requested",
    "deviation-approved", "closed",
]

# Remediation response lifecycle values (OSCAL risk.response.lifecycle)
REMEDIATION_LIFECYCLES = ["recommendation", "planned", "completed"]

# Finding target status state values (OSCAL finding.target.status.state)
FINDING_STATUS_STATES = ["satisfied", "not-satisfied"]

# Finding target status reason values (OSCAL finding.target.status.reason)
FINDING_STATUS_REASONS = ["", "pass", "fail", "other"]

# Valid component type values from OSCAL 1.2.2 schema
COMPONENT_TYPES = [
    "defined-system", "system", "interconnection", "software", "hardware",
    "service", "policy", "process", "procedure", "plan", "guidance",
    "standard", "validation", "physical",
]

# =============================================================================
# OSCAL CATALOG PARSING HELPERS
# These small functions extract specific pieces of data from the raw JSON
# dictionaries that come out of an OSCAL catalog file.
# =============================================================================

def get_prop(props, name, default="—"):
    """
    Search a list of OSCAL 'props' (properties) for one with a matching name,
    and return its value.

    In OSCAL, many objects carry extra metadata as a list of
    {"name": "...", "value": "..."} dictionaries called 'props'.
    This function makes it easy to look one up by name.

    Parameters:
        props   - A list of property dictionaries, e.g.
                  [{"name": "label", "value": "GOV-01"}, ...]
        name    - The property name to search for, e.g. "label"
        default - What to return if the property is not found (defaults to "—")

    Returns:
        The value string if found, otherwise the default.

    Example:
        props = [{"name": "label", "value": "GOV-01"}]
        get_prop(props, "label")   # returns "GOV-01"
        get_prop(props, "colour")  # returns "—"
    """
    for p in props:
        # Check whether this property's name matches what we are looking for
        if p.get("name") == name:
            return p.get("value", default)
    # If we went through every property and found nothing, return the default
    return default


def get_all_props(props, name):
    """
    Like get_prop, but returns ALL values whose name matches, as a list.

    This is needed for properties that can appear more than once —
    for example, 'applicability' (NC, OS, P, S, TS) appears as
    several separate props on one control.

    Parameters:
        props - A list of property dictionaries
        name  - The property name to search for

    Returns:
        A list of matching value strings (may be empty).

    Example:
        props = [
            {"name": "applicability", "value": "NC"},
            {"name": "applicability", "value": "OS"},
        ]
        get_all_props(props, "applicability")  # returns ["NC", "OS"]
    """
    # This is a "list comprehension" — a compact way to build a list
    # by looping and filtering in one line.
    return [p.get("value", "") for p in props if p.get("name") == name]


def get_statement(parts):
    """
    Extract the plain-English description (called the 'statement') from
    an OSCAL control's 'parts' list.

    Each control can have multiple 'parts' (e.g. statement, guidance,
    references). We only want the one named 'statement'.

    Parameters:
        parts - A list of part dictionaries from a control

    Returns:
        The prose text of the statement, or an empty string if not found.

    Example:
        parts = [{"name": "statement", "prose": "Cables are run in ..."}]
        get_statement(parts)  # returns "Cables are run in ..."
    """
    for part in parts:
        if part.get("name") == "statement":
            # .strip() removes any leading/trailing whitespace
            return part.get("prose", "").strip()
    return ""


def collect_controls(obj, parent_titles=None):
    """
    Recursively walk an OSCAL catalog (or group within a catalog) and
    collect every control into a flat list.

    OSCAL catalogs are nested: a catalog contains groups, groups can
    contain sub-groups, and groups/sub-groups contain controls.
    This function uses recursion (calling itself) to walk all levels.

    Each collected control becomes a plain Python dictionary with
    all the fields the GUI needs to display it.

    Parameters:
        obj           - A catalog, group, or sub-group dictionary
        parent_titles - A list of ancestor group titles, used to build
                        the breadcrumb path shown in the detail panel.
                        Starts as None (empty) at the top level.

    Returns:
        A flat list of control dictionaries.
    """
    # On the very first call, parent_titles is None — start with an empty list.
    # We avoid using [] as a default argument because Python reuses the same
    # list object across calls, which causes bugs. None is the safe pattern.
    if parent_titles is None:
        parent_titles = []

    result = []

    # Build the breadcrumb path for this level by adding the current title
    title = obj.get("title", "")
    path = parent_titles + ([title] if title else [])

    # Process every control directly inside this object
    for ctrl in obj.get("controls", []):
        props   = ctrl.get("props", [])
        ctrl_id = ctrl.get("id", "")

        # Use the 'label' prop (e.g. "GOV-01") as the display label.
        # If there is no label prop, fall back to the raw id (e.g. "ism-1130").
        # The 'or' operator returns ctrl_id if get_prop returns None.
        label = get_prop(props, "label", default=None) or ctrl_id

        # Build a clean dictionary with just the fields we need
        result.append({
            "id":              ctrl_id,
            "label":           label,
            "class":           ctrl.get("class", ""),    # e.g. "ISM-control"
            "title":           ctrl.get("title", ""),
            "statement":       get_statement(ctrl.get("parts", [])),
            "applicability":   get_all_props(props, "applicability"),
            "revision":        get_prop(props, "revision"),
            "updated":         get_prop(props, "updated"),
            "essential_eight": get_prop(props, "essential-eight-applicability"),
            # " › ".join([...]) builds "Group › Sub-group › ..." breadcrumb text
            "path":            " › ".join(path),
            # Top-level group title — e.g. "Guidelines for cyber security roles".
            # This is the broadest category the control belongs to.
            "guideline":       path[0] if path else "",
        })

    # Recurse into any sub-groups, passing the current path down
    for grp in obj.get("groups", []):
        result.extend(collect_controls(grp, path))

    return result


# =============================================================================
# FILE LOADERS
# These functions open a JSON file from disk, validate it, and return
# a clean Python dictionary the rest of the app can use.
# =============================================================================

def load_catalog(filepath):
    """
    Open an OSCAL catalog JSON file and return its contents as a
    structured Python dictionary.

    Parameters:
        filepath - The full path to the .json catalog file (string or Path)

    Returns:
        A dictionary containing:
            title, published, last_modified, version, oscal_version,
            controls (flat list), filepath

    Raises:
        ValueError - if the file does not look like an OSCAL catalog
        json.JSONDecodeError - if the file is not valid JSON
    """
    # Open the file and parse its JSON content into a Python dictionary.
    # 'encoding="utf-8"' ensures special characters are read correctly.
    with open(filepath, encoding="utf-8") as f:
        data = json.load(f)

    # Every OSCAL catalog file must have a top-level "catalog" key.
    # If it is missing, this is not a valid catalog file.
    if "catalog" not in data:
        raise ValueError("Missing 'catalog' key — not an OSCAL catalog.")

    catalog = data["catalog"]
    # 'metadata' holds the document title, version, dates, etc.
    meta = catalog.get("metadata", {})

    # Return a clean, flat dictionary — the GUI only needs these fields.
    # .get("key", "—") returns "—" if the key does not exist,
    # so labels always show something rather than crashing.
    return {
        "title":         meta.get("title", "Untitled catalog"),
        "published":     meta.get("published", "—"),
        "last_modified": meta.get("last-modified", "—"),
        "version":       meta.get("version", "—"),
        "oscal_version": meta.get("oscal-version", "—"),
        # collect_controls walks the whole catalog and returns a flat list
        "controls":      collect_controls(catalog),
        # Store the file path so we can reference it later when building the SSP
        "filepath":      filepath,
    }


def load_profile(filepath):
    """
    Open an OSCAL profile JSON file and return its contents as a
    structured Python dictionary.

    A profile is a document that selects a subset of controls from a
    catalog — for example, only the controls relevant to a non-classified
    system. It lists the selected control IDs under 'imports'.

    Parameters:
        filepath - The full path to the .json profile file

    Returns:
        A dictionary containing:
            title, published, last_modified, version, oscal_version,
            ids (a Python set of selected control ID strings), filepath

    Raises:
        ValueError - if the file does not look like an OSCAL profile
    """
    with open(filepath, encoding="utf-8") as f:
        data = json.load(f)

    if "profile" not in data:
        raise ValueError("Missing 'profile' key — not an OSCAL profile.")

    profile = data["profile"]
    meta    = profile.get("metadata", {})

    # Collect every selected control ID into a Python 'set'.
    # A set is like a list but automatically removes duplicates and
    # makes lookups (the 'in' operator) very fast — important when
    # filtering thousands of controls.
    included_ids = set()

    # The profile may import from multiple sources, each with multiple
    # include-controls selectors, each with multiple with-ids entries.
    # Three nested loops handle all combinations.
    for imp in profile.get("imports", []):
        for selector in imp.get("include-controls", []):
            for ctrl_id in selector.get("with-ids", []):
                included_ids.add(ctrl_id)

    return {
        "title":         meta.get("title", "Untitled profile"),
        "published":     meta.get("published", "—"),
        "last_modified": meta.get("last-modified", "—"),
        "version":       meta.get("version", "—"),
        "oscal_version": meta.get("oscal-version", "—"),
        "ids":           included_ids,
        "filepath":      filepath,
    }


# =============================================================================
# SSP DATA MODEL
# The System Security Plan (SSP) is built up from user input in the GUI.
# Internally we store everything as a plain Python dictionary (self._ssp).
# These functions create, validate, convert, and parse that dictionary.
# =============================================================================

def empty_ssp():
    """
    Return a brand-new, blank SSP dictionary with sensible default values.

    This is called when the user clicks 'New SSP' or when the app starts.
    Every key here corresponds to a field in the SSP Editor form.

    Returns:
        A dictionary with all SSP fields set to empty strings or defaults.
    """
    return {
        # A UUID is a random unique identifier — required by the OSCAL schema.
        # We generate one now so it stays stable for the life of this SSP.
        "uuid": new_uuid(),

        # ── Section 1: Metadata ───────────────────────────────────────────
        "title":            "",     # SSP document title
        "version":          "1.0",  # Document version
        "date_authorized":  "",     # Date the system was authorised (YYYY-MM-DD)

        # ── Section 2: System characteristics ────────────────────────────
        "system_name":                "",
        "system_name_short":          "",
        "system_description":         "",
        # FIPS 199 defines three impact levels: low, moderate, high
        "security_sensitivity_level": "fips-199-moderate",
        "status":                     "under-development",
        "status_remarks":             "",

        # ── OSCAL 1.2.x structured security-impact-level (M2 fix) ────────
        # OSCAL 1.2.x uses a separate field for each CIA dimension. We store
        # them individually so the UI can show three distinct dropdowns.
        "confidentiality_impact": "fips-199-moderate",
        "integrity_impact":       "fips-199-moderate",
        "availability_impact":    "fips-199-moderate",

        # ── Stable UUID for the auto-generated "this-system" component (M3 fix)
        # Left blank on first creation — generated and stored on first save so
        # it stays the same across subsequent saves of the same document.
        "this_system_component_uuid": "",

        # ── Phone field for round-trip compatibility (L3 fix) ─────────────
        # The DOCX export table has a "Phone" column. We default it to empty
        # so that old saved files (which have no phone key) still load cleanly.
        "phone": "",

        # ── Section 3: Authorization boundary ────────────────────────────
        "auth_boundary_description":  "",

        # ── Section 4: Network architecture & data flow (optional) ───────
        "network_architecture": "",
        "data_flow":            "",
        "data_flow_diagrams":   [],   # list of {uuid, caption, link, description}

        # ── Sections 5-7: Lists — start empty, user adds entries ─────────
        # Each role is: {"role_id": str, "title": str}
        "roles":            [],
        # Each party is:  {"uuid": str, "type": str, "name": str, "email": str}
        "parties":          [],
        # Each info type: {"uuid", "title", "description", "c_impact", "i_impact", "a_impact",
        #                   "component_flows": [{"component_uuid", "component_title", "direction"}]}
        "information_types": [],
        # Section 7b: responsible-parties — maps each role to the party UUID
        # who fills it.  Each entry: {"role_id": str, "party_uuid": str,
        #                              "party_name": str, "remarks": str}
        "responsible_parties": [],

        # ── Section 8: System Components ──────────────────────────────────
        # Each component is a dict with keys: uuid, type, title, description,
        # purpose, status, status_remarks, responsible_roles (list of role-id
        # strings), remarks. These become extra entries (alongside the
        # auto-generated "this-system" component) in system-implementation.
        "components":           [],

        # ── Section 9: Control Implementations ───────────────────────────
        # Each entry describes how the system implements one control:
        #   {"control_id", "remarks", "by_components": [
        #       {"uuid", "component_uuid", "description", "impl_status", "remarks"}
        #   ]}
        # These become the implemented-requirements of control-implementation.
        "ctrl_implementations": [],

        # ── Section 9b: Control-implementation set-parameters ─────────────
        # Override catalog parameter values at the SSP level.
        # Each entry: {"param_id": str, "values": [str], "remarks": str}
        "set_parameters": [],

        # ── Section 11: System Users ──────────────────────────────────────
        # People or entities that interact with the system.
        # Each entry: {"uuid", "title", "short_name", "description",
        #              "role_ids": [str], "remarks": str}
        "users": [],

        # ── Section 12: Inventory Items ───────────────────────────────────
        # Hardware/software assets that make up the system.
        # Each entry: {"uuid", "description", "props": [{"name","value"}],
        #              "implemented_components": [comp_uuid],
        #              "remarks": str}
        "inventory_items": [],
    }


def build_oscal_ssp(ssp, profile, catalog, oscal_version=None):
    """
    Convert our internal SSP dictionary into a fully valid OSCAL SSP
    JSON document (as a Python dictionary ready to be written to a file).

    The OSCAL SSP schema defines a specific structure. This function
    maps our simple internal fields to that structure.

    Parameters:
        ssp           - The internal SSP dictionary (from the form)
        profile       - The loaded profile dictionary, or None
        catalog       - The loaded catalog dictionary, or None
        oscal_version - The OSCAL version string to write into the file
                        (e.g. "1.2.2"). Must be supplied explicitly — callers
                        should pass get_oscal_version() from the toolbar.

    Returns:
        A nested dictionary matching the OSCAL system-security-plan schema.
    """
    if oscal_version is None:
        raise ValueError(
            "oscal_version must be provided explicitly (e.g. '1.2.2'). "
            "Pass get_oscal_version() from the toolbar callback."
        )
    # Record the exact moment of saving — required by the schema
    now = now_iso()

    # ── Convert roles to OSCAL format ────────────────────────────────────────
    # Our internal format:  {"role_id": "isso", "title": "ISSO"}
    # OSCAL format:         {"id": "isso", "title": "ISSO"}
    roles = [
        {"id": r["role_id"], "title": r["title"]}
        for r in ssp.get("roles", [])
    ]

    # ── Convert parties to OSCAL format ──────────────────────────────────────
    parties = []
    for p in ssp.get("parties", []):
        # Build the required fields first
        entry = {
            "uuid": p["uuid"],
            "type": p["type"],   # "person" or "organization"
            "name": p["name"],
        }
        # Email is optional — only add it if the user filled it in
        if p.get("email"):
            entry["email-addresses"] = [p["email"]]
        parties.append(entry)

    # ── Convert responsible-parties to OSCAL format ───────────────────────────
    # OSCAL responsible-parties groups all party UUIDs for the same role into
    # a single entry with a party-uuids list.  Our internal format stores one
    # row per (role, party) pair, so we aggregate here.
    # OSCAL format: [{"role-id": "isso", "party-uuids": ["uuid-1", ...]}, ...]
    rp_by_role = {}
    for rp in ssp.get("responsible_parties", []):
        role_id    = rp.get("role_id", "").strip()
        party_uuid = rp.get("party_uuid", "").strip()
        if role_id and party_uuid:
            rp_by_role.setdefault(role_id, {"role-id": role_id, "party-uuids": []})
            if party_uuid not in rp_by_role[role_id]["party-uuids"]:
                rp_by_role[role_id]["party-uuids"].append(party_uuid)
    responsible_parties = list(rp_by_role.values())

    # ── Convert information types to OSCAL format ─────────────────────────────
    # CIA = Confidentiality, Integrity, Availability — the three security pillars
    info_types = []
    for it in ssp.get("information_types", []):
        # Build component-flow props (stored as OSCAL props with class="data-flow")
        flow_props = []
        for flow in it.get("component_flows", []):
            flow_props.append({"name": "component-uuid",  "value": flow["component_uuid"],  "class": "data-flow"})
            flow_props.append({"name": "flow-direction",   "value": flow["direction"],       "class": "data-flow"})
            flow_props.append({"name": "component-title",  "value": flow["component_title"], "class": "data-flow"})
        info_types.append({
            "uuid":        it["uuid"],
            "title":       it["title"],
            "description": it.get("description", ""),
            # Each impact level is wrapped in a {"base": "fips-199-..."} dict
            "confidentiality-impact": {"base": it.get("c_impact", "fips-199-moderate")},
            "integrity-impact":       {"base": it.get("i_impact", "fips-199-moderate")},
            "availability-impact":    {"base": it.get("a_impact", "fips-199-moderate")},
            **({"props": flow_props} if flow_props else {}),
        })

    # ── Build the import-profile reference and back-matter resource ───────────
    #
    # The OSCAL SSP must declare which profile (or catalog) it is based on.
    # The recommended way is:
    #   1. Give the profile a UUID and store its details in "back-matter"
    #   2. Set import-profile.href to "#<that-uuid>" (the # means "look in this file")
    #
    # This keeps the SSP self-documenting: anyone reading the file can see
    # exactly which profile version was used, without needing the profile file.

    # Re-use the existing resource UUID if this SSP was previously saved,
    # so the internal reference stays stable across edits.
    profile_resource_uuid = ssp.get("profile_resource_uuid") or new_uuid()

    if profile:
        # A profile is loaded — reference it (preferred by the OSCAL schema)
        fp    = profile.get("filepath", "")
        fname = Path(fp).name if fp else "profile.json"  # just the filename, not full path
        import_href = f"#{profile_resource_uuid}"         # e.g. "#abc-123"
        back_matter_resource = {
            "uuid":  profile_resource_uuid,
            "title": profile.get("title", "OSCAL Profile"),
            # Store key metadata so the SSP is self-documenting
            "props": [
                {"name": "type",          "value": "profile"},
                {"name": "version",       "value": profile.get("version", "—")},
                {"name": "oscal-version", "value": profile.get("oscal_version", "—")},
                {"name": "last-modified", "value": profile.get("last_modified", "—")},
            ],
            # rlinks = relative links — points to the actual profile file
            "rlinks": [{"href": fname}],
        }
    elif catalog:
        # No profile loaded — fall back to referencing the catalog directly.
        # The OSCAL schema discourages this but it keeps the file valid.
        fp    = catalog.get("filepath", "")
        fname = Path(fp).name if fp else "catalog.json"
        import_href = f"#{profile_resource_uuid}"
        back_matter_resource = {
            "uuid":  profile_resource_uuid,
            "title": catalog.get("title", "OSCAL Catalog"),
            "props": [
                {"name": "type",          "value": "catalog"},
                {"name": "version",       "value": catalog.get("version", "—")},
                {"name": "oscal-version", "value": catalog.get("oscal_version", "—")},
                {"name": "last-modified", "value": catalog.get("last_modified", "—")},
            ],
            "rlinks": [{"href": fname}],
        }
    else:
        # Neither loaded — use a placeholder string (file will be invalid)
        import_href          = "PROFILE_OR_CATALOG_HREF"
        back_matter_resource = None

    # ── Stable UUID for the "this-system" component (M3 fix) ─────────────────
    # OSCAL requires at least one component in system-implementation.
    # We reuse the stored UUID if available so that re-saving the same SSP
    # does not generate a new UUID (which would break any references to it).
    component_uuid = ssp.get("this_system_component_uuid") or new_uuid()
    # Write the UUID back so it persists into subsequent saves via _collect()
    ssp["this_system_component_uuid"] = component_uuid

    # ── Build the full system-implementation components list ──────────────────
    # Always start with the auto-generated "this-system" component, then append
    # each user-defined component from Section 8. Optional keys (purpose,
    # status remarks, responsible-roles, remarks) are only included when present
    # using the **({...} if cond else {}) idiom described above.
    si_components = [{
        "uuid":        component_uuid,
        "type":        "this-system",
        "title":       ssp.get("system_name", "This System"),
        "description": ssp.get("system_description", ""),
        "status":      {"state": ssp.get("status", "under-development")},
    }]
    for comp in ssp.get("components", []):
        # Build the OSCAL protocols list from the internal port_ranges format.
        # Each protocol gets a fresh UUID; port start/end are written as integers.
        oscal_protocols = []
        for proto in comp.get("protocols", []):
            p = {"uuid": new_uuid(), "name": proto["name"]}
            if proto.get("title"):
                p["title"] = proto["title"]
            prs = [
                {k: v for k, v in {
                    "start":     int(pr["start"]) if pr.get("start") is not None else None,
                    "end":       int(pr["end"])   if pr.get("end")   is not None else None,
                    "transport": pr.get("transport") or None,
                    "remarks":   pr.get("remarks")   or None,
                }.items() if v is not None}
                for pr in proto.get("port_ranges", [])
            ]
            if prs:
                p["port-ranges"] = prs
            oscal_protocols.append(p)

        si_components.append({
            "uuid":        comp["uuid"],
            "type":        comp["type"],
            "title":       comp["title"],
            "description": comp["description"],
            **({"purpose": comp["purpose"]} if comp.get("purpose") else {}),
            "status": {
                # Guard against an empty string — the schema enum requires one
                # of the four valid values; default to "operational" if unset.
                "state": comp.get("status") or "operational",
                **({"remarks": comp["status_remarks"]}
                   if comp.get("status_remarks") else {}),
            },
            **({"responsible-roles": [{"role-id": r}
                                      for r in comp["responsible_roles"]]}
               if comp.get("responsible_roles") else {}),
            **({"protocols": oscal_protocols} if oscal_protocols else {}),
            **({"remarks": comp["remarks"]} if comp.get("remarks") else {}),
        })

    # ── Build implemented-requirements from Section 9 ─────────────────────────
    # Each ctrl_implementation maps to one implemented-requirement, with its
    # by_components flattened into OSCAL by-component entries. If the user has
    # not added any control implementations, we leave an empty list so the
    # control-implementation block stays schema-valid.
    implemented_requirements = [
        {
            # Use a stable UUID if this control-implementation was previously saved;
            # otherwise generate a new one. This prevents UUID churn on re-save. (M3)
            "uuid":       ci.get("uuid") or new_uuid(),
            "control-id": ci["control_id"],
            "by-components": [
                {
                    "component-uuid": bc["component_uuid"],
                    "uuid":           bc["uuid"],
                    "description":    bc["description"],
                    **({"implementation-status": {"state": bc["impl_status"]}}
                       if bc.get("impl_status") else {}),
                    **({"remarks": bc["remarks"]} if bc.get("remarks") else {}),
                }
                for bc in ci["by_components"]
            ],
            **({"remarks": ci["remarks"]} if ci.get("remarks") else {}),
        }
        for ci in ssp.get("ctrl_implementations", [])
    ]

    # ── Assemble the final OSCAL document ────────────────────────────────────
    # The ** (double-star) operator unpacks a dictionary into keyword arguments.
    # Here we use it to conditionally include optional keys only when they
    # have a value — e.g. **({"key": val} if val else {}) adds "key" only
    # when val is not empty.
    doc = {
        "system-security-plan": {
            "uuid": ssp["uuid"],

            # Metadata section — document-level information
            "metadata": {
                "title":         ssp["title"],
                "last-modified": now,
                "version":       ssp.get("version", "1.0"),
                "oscal-version": oscal_version,     # OSCAL schema version we target
                **({"roles":                roles}                if roles                else {}),
                **({"parties":             parties}             if parties             else {}),
                **({"responsible-parties": responsible_parties} if responsible_parties else {}),
            },

            # Declare which profile this SSP is based on
            "import-profile": {"href": import_href},

            # System characteristics — details about the system being described
            "system-characteristics": {
                # system-ids: a unique identifier for the system itself
                "system-ids": [{
                    "id":              ssp["uuid"],
                    "identifier-type": "https://ietf.org/rfc/rfc4122",  # UUID standard
                }],
                "system-name": ssp.get("system_name", ""),
                # Only include short name if the user provided one
                **({"system-name-short": ssp["system_name_short"]}
                   if ssp.get("system_name_short") else {}),
                "description": ssp.get("system_description", ""),
                **({"date-authorized": ssp["date_authorized"]}
                   if ssp.get("date_authorized") else {}),
                **({"security-sensitivity-level": ssp["security_sensitivity_level"]}
                   if ssp.get("security_sensitivity_level") else {}),
                # OSCAL 1.2.x adds a structured security-impact-level alongside
                # the flat sensitivity level. We emit both for compatibility. (M2)
                "security-impact-level": {
                    "security-objective-confidentiality": ssp.get(
                        "confidentiality_impact", "fips-199-moderate"),
                    "security-objective-integrity":       ssp.get(
                        "integrity_impact",       "fips-199-moderate"),
                    "security-objective-availability":    ssp.get(
                        "availability_impact",    "fips-199-moderate"),
                },
                "system-information": {
                    "information-types": info_types,
                },
                "status": {
                    "state": ssp.get("status", "under-development"),
                    **({"remarks": ssp["status_remarks"]}
                       if ssp.get("status_remarks") else {}),
                },
                "authorization-boundary": {
                    "description": ssp.get("auth_boundary_description", ""),
                },
                # Network architecture and data flow are optional
                **({"network-architecture": {"description": ssp["network_architecture"]}}
                   if ssp.get("network_architecture") else {}),
                **({"data-flow": {
                    "description": ssp["data_flow"],
                    **({"diagrams": [
                        {
                            "uuid": d["uuid"],
                            **({"description": d["description"]} if d.get("description") else {}),
                            "caption": d["caption"],
                            "links": [{"href": d["link"], "rel": "diagram"}],
                        }
                        for d in ssp.get("data_flow_diagrams", [])
                    ]} if ssp.get("data_flow_diagrams") else {}),
                }} if (ssp.get("data_flow") or ssp.get("data_flow_diagrams")) else {}),
            },

            # System implementation — the components that make up the system.
            # Includes the auto-generated "this-system" component plus any
            # user-defined components from Section 8.
            "system-implementation": {
                **({"users": [
                    {
                        "uuid":  u["uuid"],
                        "title": u["title"],
                        **({"short-name":   u["short_name"]}  if u.get("short_name")  else {}),
                        **({"description":  u["description"]} if u.get("description") else {}),
                        **({"role-ids":     u["role_ids"]}    if u.get("role_ids")    else {}),
                        **({"remarks":      u["remarks"]}     if u.get("remarks")     else {}),
                    }
                    for u in ssp.get("users", [])
                ]} if ssp.get("users") else {}),
                "components": si_components,
                **({"inventory-items": [
                    {
                        "uuid":        ii["uuid"],
                        "description": ii["description"],
                        **({"props": [
                            {"name": p["name"], "value": p["value"]}
                            for p in ii.get("props", [])
                        ]} if ii.get("props") else {}),
                        **({"implemented-components": [
                            {"component-uuid": c_uuid}
                            for c_uuid in ii.get("implemented_components", [])
                        ]} if ii.get("implemented_components") else {}),
                        **({"remarks": ii["remarks"]} if ii.get("remarks") else {}),
                    }
                    for ii in ssp.get("inventory_items", [])
                ]} if ssp.get("inventory_items") else {}),
            },

            # Control implementation — how the system's components implement
            # each control (Section 9). The description is always present;
            # implemented-requirements is empty until the user adds entries.
            "control-implementation": {
                "description": "Control implementation statements for this system.",
                **({"set-parameters": [
                    {
                        "param-id": sp["param_id"],
                        "values":   sp["values"],
                        **({"remarks": sp["remarks"]} if sp.get("remarks") else {}),
                    }
                    for sp in ssp.get("set_parameters", [])
                ]} if ssp.get("set_parameters") else {}),
                "implemented-requirements": implemented_requirements,
            },

            # Back-matter — reference documents (our profile/catalog entry goes here)
            **({"back-matter": {"resources": [back_matter_resource]}}
               if back_matter_resource else {}),
        }
    }
    return doc


def parse_ssp_file(data):
    """
    Read a previously saved OSCAL SSP JSON document (already loaded into
    a Python dictionary) and convert it back into our internal SSP format.

    This is the reverse of build_oscal_ssp — it lets the user re-open
    a saved SSP to continue editing.

    Parameters:
        data - A dictionary from json.load() containing a saved OSCAL SSP

    Returns:
        A tuple of (ssp_dict, back_matter_info):
            ssp_dict        - Our internal SSP dictionary, ready to populate the form
            back_matter_info - Details about the referenced profile/catalog,
                               extracted from the SSP's back-matter section
    """
    # Navigate into the nested OSCAL structure
    # .get("key", {}) returns an empty dict if the key is missing,
    # which lets us safely call .get() on the result without crashing.
    root   = data.get("system-security-plan", {})
    meta   = root.get("metadata", {})
    sc     = root.get("system-characteristics", {})
    status = sc.get("status", {})
    ab     = sc.get("authorization-boundary", {})
    na     = sc.get("network-architecture", {})
    df     = sc.get("data-flow", {})

    # ── Roles: convert from OSCAL format back to our internal format ──────────
    # OSCAL uses "id", we use "role_id"
    roles = [
        {"role_id": r.get("id", ""), "title": r.get("title", "")}
        for r in meta.get("roles", [])
    ]

    # ── Parties ───────────────────────────────────────────────────────────────
    parties = []
    for p in meta.get("parties", []):
        emails = p.get("email-addresses", [])
        parties.append({
            "uuid":  p.get("uuid", new_uuid()),
            "type":  p.get("type", "person"),
            "name":  p.get("name", ""),
            # Take only the first email address if there are multiple
            "email": emails[0] if emails else "",
        })

    # ── Responsible parties ───────────────────────────────────────────────────
    # OSCAL stores one entry per role with a list of party UUIDs.
    # We expand to one row per (role, party) pair so the treeview UI is flat.
    # We also look up the party name for display purposes.
    party_by_uuid = {p["uuid"]: p["name"] for p in parties}
    responsible_parties = []
    for rp in meta.get("responsible-parties", []):
        role_id = rp.get("role-id", "")
        for party_uuid in rp.get("party-uuids", []):
            responsible_parties.append({
                "role_id":    role_id,
                "party_uuid": party_uuid,
                "party_name": party_by_uuid.get(party_uuid, party_uuid),
                "remarks":    rp.get("remarks", ""),
            })

    # ── Information types ─────────────────────────────────────────────────────
    info_types = []
    for it in sc.get("system-information", {}).get("information-types", []):
        # Parse component_flows back from props (groups of 3 with class="data-flow")
        component_flows = []
        props = [p for p in it.get("props", []) if p.get("class") == "data-flow"]
        i = 0
        while i + 2 < len(props):
            if props[i].get("name") == "component-uuid":
                component_flows.append({
                    "component_uuid":  props[i]["value"],
                    "direction":       props[i+1]["value"] if i+1 < len(props) else "internal",
                    "component_title": props[i+2]["value"] if i+2 < len(props) else "",
                })
                i += 3
            else:
                i += 1
        info_types.append({
            "uuid":        it.get("uuid", new_uuid()),
            "title":       it.get("title", ""),
            "description": it.get("description", ""),
            # Unpack the nested {"base": "fips-199-..."} structure
            "c_impact": it.get("confidentiality-impact", {}).get("base", "fips-199-moderate"),
            "i_impact": it.get("integrity-impact",       {}).get("base", "fips-199-moderate"),
            "a_impact": it.get("availability-impact",    {}).get("base", "fips-199-moderate"),
            "component_flows": component_flows,
        })

    # ── System components (Section 8) ─────────────────────────────────────────
    # Read every component except the auto-generated "this-system" one, which
    # we regenerate from the system characteristics when saving (so re-importing
    # it would create a duplicate).
    components = []
    for c in root.get("system-implementation", {}).get("components", []):
        if c.get("type") == "this-system":
            continue
        status_obj  = c.get("status", {})
        comp_roles  = [r.get("role-id", "") for r in c.get("responsible-roles", [])]
        # Parse protocols back into the internal port_ranges format
        protocols = []
        for proto in c.get("protocols", []):
            prs = [
                {
                    "start":     pr.get("start", 0),
                    "end":       pr.get("end", pr.get("start", 0)),
                    "transport": pr.get("transport", "TCP"),
                    "remarks":   pr.get("remarks", ""),
                }
                for pr in proto.get("port-ranges", [])
            ]
            protocols.append({
                "name":        proto.get("name", ""),
                "title":       proto.get("title", ""),
                "port_ranges": prs,
            })
        components.append({
            "uuid":              c.get("uuid", new_uuid()),
            "type":              c.get("type", "software"),
            "title":             c.get("title", ""),
            "description":       c.get("description", ""),
            "purpose":           c.get("purpose", ""),
            "status":            status_obj.get("state", "operational"),
            "status_remarks":    status_obj.get("remarks", ""),
            "responsible_roles": [r for r in comp_roles if r],
            "protocols":         protocols,
            "remarks":           c.get("remarks", ""),
        })

    # ── System users (Section 11) ─────────────────────────────────────────────
    users = []
    for u in root.get("system-implementation", {}).get("users", []):
        users.append({
            "uuid":        u.get("uuid", new_uuid()),
            "title":       u.get("title", ""),
            "short_name":  u.get("short-name", ""),
            "description": u.get("description", ""),
            "role_ids":    u.get("role-ids", []),
            "remarks":     u.get("remarks", ""),
        })

    # ── Inventory items (Section 12) ──────────────────────────────────────────
    inventory_items = []
    for ii in root.get("system-implementation", {}).get("inventory-items", []):
        inventory_items.append({
            "uuid":        ii.get("uuid", new_uuid()),
            "description": ii.get("description", ""),
            "props":       [
                {"name": p.get("name", ""), "value": p.get("value", "")}
                for p in ii.get("props", [])
            ],
            "implemented_components": [
                ic.get("component-uuid", "")
                for ic in ii.get("implemented-components", [])
                if ic.get("component-uuid")
            ],
            "remarks": ii.get("remarks", ""),
        })

    # ── Control-implementation set-parameters (Section 9b) ────────────────────
    set_parameters = []
    for sp in root.get("control-implementation", {}).get("set-parameters", []):
        set_parameters.append({
            "param_id": sp.get("param-id", ""),
            "values":   sp.get("values", []),
            "remarks":  sp.get("remarks", ""),
        })

    # ── Control implementations (Section 9) ───────────────────────────────────
    # Group the OSCAL implemented-requirements by control-id. A well-formed SSP
    # has one implemented-requirement per control, but we deduplicate defensively
    # so multiple entries for the same control merge their by-components.
    ctrl_implementations = []
    seen_ctrl_ids = {}   # control_id -> the entry dict already in the list
    for ir in root.get("control-implementation", {}).get("implemented-requirements", []):
        ctrl_id = ir.get("control-id", "")
        if not ctrl_id:
            continue
        bcs = []
        for bc in ir.get("by-components", []):
            is_obj = bc.get("implementation-status", {})
            bcs.append({
                "uuid":           bc.get("uuid", new_uuid()),
                "component_uuid": bc.get("component-uuid", ""),
                "description":    bc.get("description", ""),
                "impl_status":    is_obj.get("state", "implemented"),
                "remarks":        bc.get("remarks", ""),
            })
        if ctrl_id in seen_ctrl_ids:
            seen_ctrl_ids[ctrl_id]["by_components"].extend(bcs)
        else:
            entry = {
                "control_id":    ctrl_id,
                "remarks":       ir.get("remarks", ""),
                "by_components": bcs,
            }
            seen_ctrl_ids[ctrl_id] = entry
            ctrl_implementations.append(entry)

    # ── Data flow diagrams ────────────────────────────────────────────────────
    data_flow_diagrams = []
    for d in df.get("diagrams", []):
        link = ""
        for lnk in d.get("links", []):
            if lnk.get("rel") == "diagram":
                link = lnk.get("href", "")
                break
        data_flow_diagrams.append({
            "uuid":        d.get("uuid", new_uuid()),
            "caption":     d.get("caption", ""),
            "link":        link,
            "description": d.get("description", ""),
        })

    # ── Resolve import-profile reference ──────────────────────────────────────
    import_href = root.get("import-profile", {}).get("href", "")

    # If the href starts with "#", it is an internal reference to a back-matter
    # resource. We look up that resource to get the profile title and filename.
    back_matter_info = {}
    if import_href.startswith("#"):
        # Strip the "#" to get the UUID we are looking for
        ref_uuid  = import_href.lstrip("#")
        resources = root.get("back-matter", {}).get("resources", [])

        # Search the resources list for the one with a matching UUID.
        # next(..., None) returns None if nothing matches (instead of crashing).
        resource = next(
            (r for r in resources if r.get("uuid") == ref_uuid),
            None
        )
        if resource:
            rlinks = resource.get("rlinks", [])
            back_matter_info = {
                "title":   resource.get("title", ""),
                "file":    rlinks[0].get("href", "") if rlinks else "",
                # Find the "version" prop using a generator expression
                "version": next(
                    (p["value"] for p in resource.get("props", [])
                     if p.get("name") == "version"),
                    ""  # default if not found
                ),
            }
    else:
        # Bare filename — treat the href itself as the file reference
        back_matter_info = {"title": "", "file": import_href, "version": ""}

    # ── Read security-impact-level (M2 fix) ──────────────────────────────────
    # OSCAL 1.2.x stores a structured object with three separate CIA objectives.
    # We read it back so the three dropdowns are populated on re-open.
    sil = sc.get("security-impact-level", {})

    # ── Read back the stable "this-system" component UUID (M3 fix) ───────────
    # Find the component whose type is "this-system" to recover its UUID.
    this_system_uuid = ""
    for c in root.get("system-implementation", {}).get("components", []):
        if c.get("type") == "this-system":
            this_system_uuid = c.get("uuid", "")
            break

    # ── Assemble and return the internal SSP dictionary ───────────────────────
    ssp = {
        "uuid":            root.get("uuid", new_uuid()),
        "title":           meta.get("title", ""),
        "version":         meta.get("version", "1.0"),
        "date_authorized": sc.get("date-authorized", ""),
        "system_name":     sc.get("system-name", ""),
        "system_name_short": sc.get("system-name-short", ""),
        "system_description": sc.get("description", ""),
        "security_sensitivity_level": sc.get("security-sensitivity-level", "fips-199-moderate"),
        # Read back the three CIA objectives from the structured field (M2)
        "confidentiality_impact": sil.get(
            "security-objective-confidentiality", "fips-199-moderate"),
        "integrity_impact":       sil.get(
            "security-objective-integrity",       "fips-199-moderate"),
        "availability_impact":    sil.get(
            "security-objective-availability",    "fips-199-moderate"),
        "status":          status.get("state", "under-development"),
        "status_remarks":  status.get("remarks", ""),
        "auth_boundary_description": ab.get("description", ""),
        "network_architecture": na.get("description", ""),
        "data_flow":            df.get("description", ""),
        "data_flow_diagrams":   data_flow_diagrams,
        "roles":                roles,
        "parties":              parties,
        "responsible_parties":  responsible_parties,
        "information_types":    info_types,
        "components":           components,
        "ctrl_implementations": ctrl_implementations,
        "set_parameters":       set_parameters,
        "users":                users,
        "inventory_items":      inventory_items,
        # Preserve the stable "this-system" component UUID across saves (M3)
        "this_system_component_uuid": this_system_uuid,
        # Phone field — always blank on parse (OSCAL has no phone field), but
        # must exist so build_ssp_docx() can call p.get("phone", "") safely (L3)
        "phone": "",
        # Preserve the import href so it round-trips correctly on re-save
        "import_href":           import_href,
        # Preserve the back-matter UUID so it stays stable across edits
        "profile_resource_uuid": import_href.lstrip("#") if import_href.startswith("#") else "",
    }
    return ssp, back_matter_info


def validate_ssp(ssp, profile, catalog):
    """
    Check that the SSP has all the required fields before saving.

    Returns two separate lists:
      - errors:   Problems that MUST be fixed — saving is blocked.
      - warnings: Advisory issues — the user is warned but can still save.

    Parameters:
        ssp     - The internal SSP dictionary
        profile - The loaded profile dict, or None
        catalog - The loaded catalog dict, or None

    Returns:
        A tuple of (errors_list, warnings_list)
    """
    errors   = []
    warnings = []

    # ── Hard errors — required by the OSCAL schema ────────────────────────────
    if not ssp.get("title"):
        errors.append("SSP Title is required (Section 1).")
    if not ssp.get("system_name"):
        errors.append("System Name (Full) is required (Section 2).")
    if not ssp.get("system_description"):
        errors.append("System Description is required (Section 2).")
    if not ssp.get("auth_boundary_description"):
        errors.append("Authorization Boundary Description is required (Section 3).")
    if not ssp.get("information_types"):
        errors.append("At least one Information Type is required (Section 5).")
    # OSCAL 1.2.2 requires at least one system user in system-implementation.
    # An empty list or missing key both fail this check. (H4/H6 fix)
    if not ssp.get("users"):
        errors.append(
            "System Implementation: at least one System User must be defined "
            "(OSCAL requires users[])."
        )
    if not profile and not catalog:
        errors.append("No catalog or profile loaded — import-profile cannot be set.")

    # ── Soft warnings — valid but not recommended ─────────────────────────────
    if not profile:
        # Saving without a profile means we reference the catalog directly,
        # which the OSCAL spec discourages (profile is preferred).
        warnings.append(
            "No profile loaded. It is strongly recommended to load an OSCAL profile "
            "before saving — the SSP will fall back to referencing the catalog directly."
        )

    return errors, warnings


# =============================================================================
# GENERAL UTILITIES
# Small helper functions used throughout the app.
# =============================================================================

def now_iso():
    """
    Return the current date and time as an ISO 8601 string in UTC.

    Example return value: "2026-06-25T10:30:00Z"
    The OSCAL schema requires this format for 'last-modified' fields.
    """
    # timezone.utc ensures we use UTC (Coordinated Universal Time),
    # not the local time zone of the computer running the app.
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def new_uuid():
    """
    Generate and return a new random UUID (Universally Unique Identifier)
    as a string.

    Example return value: "550e8400-e29b-41d4-a716-446655440000"
    UUIDs are used throughout OSCAL to uniquely identify documents,
    components, parties, and information types.
    """
    # uuid.uuid4() generates a random UUID object; str() converts it to a string
    return str(uuid.uuid4())


def _make_oscal_validator(schema):
    """
    Build a jsonschema Draft7Validator that tolerates ECMA-262 regex patterns.

    OSCAL schemas use regex syntax (e.g. \\p{L}) that Python's re module does
    not support.  Rather than crashing or silently skipping all pattern checks,
    we wrap the 'pattern' keyword so that any pattern Python cannot compile is
    simply skipped — the field is treated as valid for that check only.

    How it works: jsonschema allows you to override how individual keywords are
    handled via validators.extend(). We replace the built-in 'pattern' handler
    with our own that catches re.error (meaning Python can't parse the regex)
    and silently returns — so the field passes that check.
    """
    # 're' is already imported at the top of this file (L2 fix — no inline import)
    def _pattern(validator, patrn, instance, schema):
        # Only check strings — other types have no regex to match against
        if not isinstance(instance, str):
            return
        try:
            if not re.search(patrn, instance):
                # 'yield' here is how jsonschema handlers report errors —
                # the validator iterates over yielded ValidationError objects
                yield jsonschema.ValidationError(
                    f"{instance!r} does not match {patrn!r}"
                )
        except re.error:
            pass  # ECMA-262 pattern Python can't compile — skip the check

    # Create a custom validator class that inherits all Draft7 rules but
    # uses our _pattern function for the 'pattern' keyword.
    OSCALValidator = jsonschema.validators.extend(
        jsonschema.Draft7Validator,
        {"pattern": _pattern},
    )
    return OSCALValidator(schema)


def validate_oscal_file(data, schema_name, zip_path):
    """
    Validate a parsed OSCAL JSON dict against the schema bundled in a zip.

    Parameters:
        data        - Parsed JSON dict (the raw file contents, not our internal
                      representation — call this before load_catalog etc.)
        schema_name - Filename inside the zip's json/schema/ folder, e.g.
                      'oscal_catalog_schema.json'
        zip_path    - pathlib.Path to the OSCAL release zip file

    Returns:
        (valid: bool, errors: list[str])
        valid  — True if the file passes schema validation (or jsonschema is
                 not installed), False if there are violations.
        errors — Human-readable error messages, one per violation (empty when
                 valid is True).  At most 10 errors are returned so the dialog
                 stays readable.
    """
    if not _JSONSCHEMA_AVAILABLE:
        # jsonschema is an optional dependency. Tell the caller validation was
        # skipped rather than silently pretending the file is valid. (L4 fix)
        return True, ["jsonschema library not installed — schema validation was skipped. "
                      "Install it with: pip install jsonschema"]

    schema_path = f"json/schema/{schema_name}"
    try:
        with zipfile.ZipFile(zip_path) as zf:
            with zf.open(schema_path) as f:
                schema = json.load(f)
    except (KeyError, zipfile.BadZipFile) as exc:
        return True, [f"Could not read schema from zip: {exc}"]

    validator = _make_oscal_validator(schema)
    errors = sorted(validator.iter_errors(data), key=lambda e: list(e.path))

    if not errors:
        return True, []

    messages = []
    for err in errors[:10]:
        location = " → ".join(str(p) for p in err.absolute_path) or "(root)"
        messages.append(f"• {location}: {err.message}")
    if len(errors) > 10:
        messages.append(f"  … and {len(errors) - 10} more error(s).")
    return False, messages


def refresh_ctrl_list(ctrl_responses, all_controls, search_term,
                      ctrl_tree, applied_tree, notebook, progress_lbl):
    """
    Rebuild the All Controls and Applied Controls Treeview tabs.

    Shared by ComponentTab, CapabilityTab, and SSPTab so the control-list
    rendering logic only exists in one place.

    Parameters:
        ctrl_responses - dict {control_id: response_text}.  A non-empty value
                         means the control has been addressed and gets a filled
                         dot (●). An empty string means no entry yet (○).
        all_controls   - list of control dicts (id, label, title, statement)
        search_term    - text filter applied to the All Controls tab only;
                         empty string means "show all"
        ctrl_tree      - Treeview widget for the All Controls tab
        applied_tree   - Treeview widget for the Applied Controls tab
        notebook       - ttk.Notebook holding both tabs (tab labels are updated
                         to show the current count, e.g. "Applied Controls (3)")
        progress_lbl   - Label widget showing "N of M controls have responses"
    """
    DOT_DONE  = "●"
    DOT_EMPTY = "○"

    term = search_term.lower().strip()
    # Build the filtered list for the All Controls tab.
    # The conditional expression "X if condition else Y" is equivalent to:
    #   if term: filtered = [list comprehension]
    #   else:    filtered = all_controls
    filtered = (
        [c for c in all_controls
         if term in c["label"].lower()
         or term in c["statement"].lower()
         or term in c["id"].lower()]
        if term else all_controls
    )

    total      = len(all_controls)
    done_count = sum(
        1 for c in all_controls
        if ctrl_responses.get(c["id"], "").strip()
    )

    def insert_row(tree, ctrl):
        has_resp = bool(ctrl_responses.get(ctrl["id"], "").strip())
        dot = DOT_DONE  if has_resp else DOT_EMPTY
        tag = "done"    if has_resp else "empty"
        # iid is the item's internal ID in the Treeview — we use the control
        # ID so we can later look up the selected control directly with
        # tree.selection()[0] instead of having to map row index → control.
        tree.insert(
            "", "end", iid=ctrl["id"],
            values=(dot, ctrl["label"], ctrl["statement"] or ctrl["title"]),
            tags=(tag,),
        )

    ctrl_tree.delete(*ctrl_tree.get_children())
    for ctrl in filtered:
        insert_row(ctrl_tree, ctrl)

    applied_tree.delete(*applied_tree.get_children())
    applied_count = 0
    for ctrl in all_controls:
        if ctrl_responses.get(ctrl["id"], "").strip():
            insert_row(applied_tree, ctrl)
            applied_count += 1

    try:
        notebook.tab(0, text=f"All Controls ({len(filtered)})")
        notebook.tab(1, text=f"Applied Controls ({applied_count})")
    except Exception:
        pass

    if total > 0:
        progress_lbl.config(text=f"{done_count} of {total} controls have responses")
    else:
        progress_lbl.config(text="")


def get_source_href(profile, catalog):
    """
    Return the filename to use as the control-implementations 'source' URI.

    The OSCAL component schema requires a 'source' field on every
    control-implementations block. It should point to the catalog or profile
    that defines the controls being implemented.

    Preference order:
      1. The profile filename (if a profile is loaded)
      2. The catalog filename (if only a catalog is loaded)
      3. A placeholder string (if neither is loaded — file will be incomplete)

    Parameters:
        profile - The loaded profile dict (with a 'filepath' key), or None
        catalog - The loaded catalog dict (with a 'filepath' key), or None

    Returns:
        A filename string, e.g. "ISM_profile.json" or "PROFILE_OR_CATALOG_HREF"
    """
    if profile and profile.get("filepath"):
        return Path(profile["filepath"]).name
    if catalog and catalog.get("filepath"):
        return Path(catalog["filepath"]).name
    return "PROFILE_OR_CATALOG_HREF"


def get_profile_controls(catalog, profile):
    """
    Return the list of controls to display in the control implementation panel.

    Used by both ComponentTab and CapabilityTab so the filtering logic lives
    in one place.

    If a profile is loaded, only controls whose IDs appear in the profile's
    baseline are returned — this is the standard OSCAL workflow where a
    profile tailors a catalog to a specific context (e.g. a classification
    level or system type).

    If no profile is loaded, all catalog controls are returned so the user
    can still create and save components without needing a profile.

    Parameters:
        catalog - The loaded catalog dict (from load_catalog()), or None
        profile - The loaded profile dict (from load_profile()), or None

    Returns:
        A list of control dicts. Empty list if no catalog is loaded.
    """
    if not catalog:
        return []
    if not profile:
        # No profile loaded — fall back to the full catalog
        return catalog["controls"]
    # Filter to only controls whose IDs are in the profile's selected set.
    # profile["ids"] is a Python set, so 'in' here is O(1) — very fast.
    return [c for c in catalog["controls"] if c["id"] in profile["ids"]]


def build_component_oscal_entry(comp, source_href):
    """
    Convert one internal component dict into an OSCAL defined-component dict.

    Used by both ComponentTab (when saving a single component file) and
    CapabilityTab (when exporting a capability document that embeds its
    member components).  Keeping the conversion in one place ensures the
    two outputs always stay in sync.

    Parameters:
        comp        - Internal component dict (uuid, type, title, description,
                      purpose, status, status_remarks, props, roles,
                      ctrl_responses, remarks).
        source_href - URI for the control source (catalog or profile href),
                      written into each control-implementations block.

    Returns:
        A dict that conforms to the OSCAL defined-component assembly.
    """
    c = {
        "uuid":        comp["uuid"],
        "type":        comp.get("type", "software"),
        "title":       comp.get("title", ""),
        "description": comp.get("description", ""),
    }

    if comp.get("purpose"):
        c["purpose"] = comp["purpose"]

    # Props — operational-status first, then any user-defined props
    props = []
    if comp.get("status"):
        p = {"name": "operational-status", "value": comp["status"]}
        if comp.get("status_remarks"):
            p["remarks"] = comp["status_remarks"]
        props.append(p)
    for prop in comp.get("props", []):
        entry = {"name": prop["name"], "value": prop["value"]}
        if prop.get("remarks"):
            entry["remarks"] = prop["remarks"]
        props.append(entry)
    if props:
        c["props"] = props

    # Responsible roles — OSCAL uses "role-id" (hyphenated)
    roles = [
        {"role-id": r["role_id"],
         **({"remarks": r["remarks"]} if r.get("remarks") else {})}
        for r in comp.get("roles", [])
    ]
    if roles:
        c["responsible-roles"] = roles

    # Protocols — network services this component exposes or uses
    protocols = []
    for proto in comp.get("protocols", []):
        p = {
            "uuid": new_uuid(),
            "name": proto["name"],
        }
        if proto.get("title"):
            p["title"] = proto["title"]
        prs = []
        for pr in proto.get("port_ranges", []):
            pr_entry = {}
            if pr.get("start") is not None:
                pr_entry["start"] = int(pr["start"])
            if pr.get("end") is not None:
                pr_entry["end"] = int(pr["end"])
            if pr.get("transport"):
                pr_entry["transport"] = pr["transport"]
            if pr.get("remarks"):
                pr_entry["remarks"] = pr["remarks"]
            if pr_entry:
                prs.append(pr_entry)
        if prs:
            p["port-ranges"] = prs
        protocols.append(p)
    if protocols:
        c["protocols"] = protocols

    # Links — external references (vendor docs, CVE advisories, policy docs, etc.)
    links = []
    for lnk in comp.get("links", []):
        href = lnk.get("href", "").strip()
        rel  = lnk.get("rel",  "").strip()
        if not href or not rel:
            continue
        entry = {"href": href, "rel": rel}
        if lnk.get("text", "").strip():
            entry["text"] = lnk["text"].strip()
        links.append(entry)
    if links:
        c["links"] = links

    # Control implementations — only include controls with a non-empty response
    implemented = [
        {
            "uuid":        new_uuid(),
            "control-id":  ctrl_id,
            "description": desc.strip(),
            "implementation-status": {
                "state": comp.get("ctrl_impl_status", {}).get(ctrl_id, "implemented")
            },
        }
        for ctrl_id, desc in comp.get("ctrl_responses", {}).items()
        if desc.strip()
    ]
    if implemented:
        c["control-implementations"] = [{
            "uuid":        new_uuid(),
            "source":      source_href,
            "description": (
                f"Control implementations for "
                f"{comp.get('title', 'this component')}."
            ),
            "implemented-requirements": implemented,
        }]

    if comp.get("remarks"):
        c["remarks"] = comp["remarks"]

    return c


# =============================================================================
# SSP WORD DOCUMENT EXPORT
# Converts the internal SSP dictionary into a formatted Microsoft Word (.docx)
# file using the python-docx library. Each SSP section becomes a Word heading
# so the document has a navigable structure and can be used as a formal report.
# =============================================================================

def build_ssp_docx(ssp, catalog=None):
    """
    Build and return a python-docx Document object from an internal SSP dict.

    The caller is responsible for saving the Document to a file path:
        doc = build_ssp_docx(ssp, catalog)
        doc.save("my_ssp.docx")

    Returns None if python-docx is not installed (caller should warn the user).

    Parameters:
        ssp     - internal SSP dict (from empty_ssp / parse_ssp_file)
        catalog - internal catalog dict (from load_catalog), or None.
                  When provided, Section 8 control implementations are sorted
                  and grouped by catalog guideline heading in catalog order.
                  When None, controls are listed in the order they were added.

    The document structure mirrors the SSP Editor sections:
        Cover — system name, version, date, status
        1. System Characteristics
        2. Authorization Boundary
        3. Network Architecture & Data Flow
        4. Information Types (table)
        5. Roles (table)
        6. Parties / People & Organisations (table)
        7. System Components (table)
        8. Control Implementations (grouped by guideline when catalog provided)
    """
    if not _DOCX_AVAILABLE:
        return None

    doc = Document()

    # ── Helper: apply a consistent style to every table ──────────────────────
    # "Table Grid" is a built-in Word style that draws borders on all cells.
    def styled_table(cols):
        t = doc.add_table(rows=0, cols=cols)
        t.style = "Table Grid"
        return t

    # ── Helper: add a header row to a table with bold, shaded cells ──────────
    def add_header_row(table, headings):
        row = table.add_row()
        for i, heading in enumerate(headings):
            cell = row.cells[i]
            cell.text = heading
            # Bold the heading text
            for run in cell.paragraphs[0].runs:
                run.bold = True
            # Light grey background using XML shading
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "D9D9D9")
            cell._tc.get_or_add_tcPr().append(shd)

    # ── Helper: add a labelled paragraph (bold label + plain value) ──────────
    def labelled(label, value):
        if not value:
            return
        p = doc.add_paragraph()
        p.add_run(label + ": ").bold = True
        p.add_run(value)

    # ────────────────────────────────────────────────────────────────────────
    # COVER
    # ────────────────────────────────────────────────────────────────────────
    title = doc.add_heading(ssp.get("system_name") or "System Security Plan", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("System Security Plan").bold = True

    doc.add_paragraph()  # blank line
    labelled("Version",         ssp.get("version", ""))
    labelled("Date",            ssp.get("date_authorized", ""))
    labelled("Status",          ssp.get("status", ""))
    labelled("System ID",       ssp.get("uuid", ""))
    labelled("Security Level",  ssp.get("security_sensitivity_level", ""))

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 1 — SYSTEM CHARACTERISTICS
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("1.  System Characteristics", level=1)
    labelled("System Name",        ssp.get("system_name", ""))
    labelled("System ID",          ssp.get("uuid", ""))
    labelled("Version",            ssp.get("version", ""))
    labelled("Date",               ssp.get("date_authorized", ""))
    labelled("Operational Status", ssp.get("status", ""))

    desc = ssp.get("system_description", "")
    if desc:
        doc.add_heading("System Description", level=2)
        doc.add_paragraph(desc)

    status_remarks = ssp.get("status_remarks", "")
    if status_remarks:
        doc.add_heading("Status Remarks", level=2)
        doc.add_paragraph(status_remarks)

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 2 — AUTHORIZATION BOUNDARY
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("2.  Authorization Boundary", level=1)
    boundary = ssp.get("auth_boundary_description", "")
    doc.add_paragraph(boundary if boundary else "(Not specified)")

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 3 — NETWORK ARCHITECTURE & DATA FLOW
    # ────────────────────────────────────────────────────────────────────────
    network = ssp.get("network_architecture", "")
    dataflow = ssp.get("data_flow", "")
    if network or dataflow:
        doc.add_heading("3.  Network Architecture & Data Flow", level=1)
        if network:
            doc.add_heading("Network Architecture", level=2)
            doc.add_paragraph(network)
        if dataflow:
            doc.add_heading("Data Flow", level=2)
            doc.add_paragraph(dataflow)

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 4 — INFORMATION TYPES
    # ────────────────────────────────────────────────────────────────────────
    info_types = ssp.get("information_types", [])
    doc.add_heading("4.  Information Types", level=1)
    if info_types:
        t = styled_table(5)
        add_header_row(t, ["Title", "Description", "Confidentiality", "Integrity", "Availability"])
        for it in info_types:
            row = t.add_row()
            row.cells[0].text = it.get("title", "")
            row.cells[1].text = it.get("description", "")
            row.cells[2].text = it.get("c_impact", "")
            row.cells[3].text = it.get("i_impact", "")
            row.cells[4].text = it.get("a_impact", "")
    else:
        doc.add_paragraph("(No information types defined)")

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 5 — ROLES
    # ────────────────────────────────────────────────────────────────────────
    roles = ssp.get("roles", [])
    doc.add_heading("5.  Roles", level=1)
    if roles:
        t = styled_table(2)
        add_header_row(t, ["Role ID", "Title"])
        for r in roles:
            row = t.add_row()
            row.cells[0].text = r.get("role_id", "")
            row.cells[1].text = r.get("title", "")
    else:
        doc.add_paragraph("(No roles defined)")

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 6 — PARTIES / PEOPLE & ORGANISATIONS
    # ────────────────────────────────────────────────────────────────────────
    parties = ssp.get("parties", [])
    doc.add_heading("6.  Parties", level=1)
    if parties:
        t = styled_table(4)
        add_header_row(t, ["Type", "Name", "Email", "Phone"])
        for p in parties:
            row = t.add_row()
            row.cells[0].text = p.get("type", "")
            row.cells[1].text = p.get("name", "")
            row.cells[2].text = p.get("email", "")
            row.cells[3].text = p.get("phone", "")
    else:
        doc.add_paragraph("(No parties defined)")

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 7 — SYSTEM USERS
    # ────────────────────────────────────────────────────────────────────────
    users = ssp.get("users", [])
    doc.add_heading("7.  System Users", level=1)
    if users:
        t = styled_table(4)
        add_header_row(t, ["Title", "Short Name", "Role IDs", "Description"])
        for u in users:
            row = t.add_row()
            row.cells[0].text = u.get("title", "")
            row.cells[1].text = u.get("short_name", "")
            row.cells[2].text = ", ".join(u.get("role_ids", []))
            row.cells[3].text = u.get("description", "")
    else:
        doc.add_paragraph("(No system users defined)")

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 7b — SYSTEM COMPONENTS
    # ────────────────────────────────────────────────────────────────────────
    components = ssp.get("components", [])
    doc.add_heading("7b.  System Components", level=1)
    if components:
        t = styled_table(4)
        add_header_row(t, ["Title", "Type", "Status", "Description"])
        for comp in components:
            row = t.add_row()
            row.cells[0].text = comp.get("title", "")
            row.cells[1].text = comp.get("type", "")
            row.cells[2].text = comp.get("status", "")
            row.cells[3].text = comp.get("description", "")
        # Add purpose/remarks as sub-paragraphs for components that have them
        for comp in components:
            if comp.get("purpose") or comp.get("remarks"):
                doc.add_heading(comp.get("title", "Component"), level=2)
                if comp.get("purpose"):
                    labelled("Purpose", comp["purpose"])
                if comp.get("responsible_roles"):
                    labelled("Responsible Roles",
                             ", ".join(comp["responsible_roles"]))
                if comp.get("remarks"):
                    labelled("Remarks", comp["remarks"])
    else:
        doc.add_paragraph("(No components defined)")

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 7c — INVENTORY ITEMS
    # ────────────────────────────────────────────────────────────────────────
    inventory_items = ssp.get("inventory_items", [])
    doc.add_heading("7c.  Inventory Items", level=1)
    if inventory_items:
        t = styled_table(3)
        add_header_row(t, ["Description", "Properties", "Components"])
        comp_titles = {c["uuid"]: c.get("title", c["uuid"]) for c in ssp.get("components", [])}
        for ii in inventory_items:
            row = t.add_row()
            row.cells[0].text = ii.get("description", "")
            props_text = "\n".join(
                f"{p['name']}: {p['value']}" for p in ii.get("props", [])
            )
            row.cells[1].text = props_text
            comp_names = ", ".join(
                comp_titles.get(c_uuid, c_uuid)
                for c_uuid in ii.get("implemented_components", [])
            )
            row.cells[2].text = comp_names
    else:
        doc.add_paragraph("(No inventory items defined)")

    # ────────────────────────────────────────────────────────────────────────
    # SECTION 8 — CONTROL IMPLEMENTATIONS
    # When a catalog is provided, controls are sorted and grouped under their
    # catalog guideline headings in catalog order. Controls in the SSP that
    # do not appear in the catalog (e.g. from a different catalog version) are
    # collected into an "Other Controls" group at the end.
    # Each control gets a Heading 2; each by-component entry is a table row.
    # ────────────────────────────────────────────────────────────────────────
    ctrl_impls = ssp.get("ctrl_implementations", [])
    doc.add_heading("8.  Control Implementations", level=1)

    set_params = ssp.get("set_parameters", [])
    if set_params:
        doc.add_heading("Parameter Overrides", level=2)
        t = styled_table(3)
        add_header_row(t, ["Parameter ID", "Values", "Remarks"])
        for sp in set_params:
            row = t.add_row()
            row.cells[0].text = sp.get("param_id", "")
            row.cells[1].text = ", ".join(sp.get("values", []))
            row.cells[2].text = sp.get("remarks", "")

    if ctrl_impls:
        # Lookup: component UUID → display title
        comp_titles = {
            c["uuid"]: c.get("title", c["uuid"])
            for c in ssp.get("components", [])
        }

        def _write_ctrl_impl(ci):
            """Write one control's heading + by-component table."""
            ctrl_id = ci.get("control_id", "")
            doc.add_heading(ctrl_id, level=2)
            if ci.get("remarks"):
                doc.add_paragraph(ci["remarks"])
            bcs = ci.get("by_components", [])
            if bcs:
                t = styled_table(3)
                add_header_row(t, ["Component", "Status", "Description"])
                for bc in bcs:
                    comp_name = comp_titles.get(
                        bc.get("component_uuid", ""), bc.get("component_uuid", "")
                    )
                    row = t.add_row()
                    row.cells[0].text = comp_name
                    row.cells[1].text = bc.get("impl_status", "")
                    row.cells[2].text = bc.get("description", "")
                    if bc.get("remarks"):
                        p = row.cells[2].add_paragraph(bc["remarks"])
                        p.runs[0].italic = True
            else:
                doc.add_paragraph("(No component responses recorded)")

        if catalog:
            # ── Catalog-ordered output ────────────────────────────────────
            # Build an index: control_id → ctrl_impl dict for fast lookup.
            impl_index = {ci["control_id"]: ci for ci in ctrl_impls}

            # Walk the catalog controls in their natural order, grouping by
            # guideline. A new Heading 1 sub-section is added each time the
            # guideline changes.
            current_guideline = None
            written_ids = set()

            for ctrl in catalog.get("controls", []):
                ctrl_id   = ctrl.get("id", "")
                guideline = ctrl.get("guideline", "")

                if ctrl_id not in impl_index:
                    continue  # This control has no SSP entry — skip it.

                # Emit a new guideline heading when the group changes.
                if guideline != current_guideline:
                    current_guideline = guideline
                    doc.add_heading(guideline or "Ungrouped Controls", level=1)

                _write_ctrl_impl(impl_index[ctrl_id])
                written_ids.add(ctrl_id)

            # Any SSP controls not found in the catalog go into a final group.
            leftover = [ci for ci in ctrl_impls
                        if ci["control_id"] not in written_ids]
            if leftover:
                doc.add_heading("Other Controls", level=1)
                for ci in leftover:
                    _write_ctrl_impl(ci)
        else:
            # ── No catalog: list controls in the order they were added ────
            for ci in ctrl_impls:
                _write_ctrl_impl(ci)
    else:
        doc.add_paragraph("(No control implementations defined)")

    return doc


# =============================================================================
# POA&M DATA FUNCTIONS
# =============================================================================

def empty_poam():
    """Return a blank POA&M working dictionary used by POAMTab."""
    return {
        "uuid":        new_uuid(),
        "title":       "",
        "version":     "1.0",
        "import_ssp":  "",   # href to the referenced SSP
        "system_id":   "",
        "observations": [],  # list of observation dicts
        "risks":        [],  # list of risk dicts
        "findings":     [],  # list of finding dicts
        "poam_items":   [],  # list of poam-item dicts
    }


def build_oscal_poam(poam, oscal_version=None, save_path=None):
    """
    Convert the internal POAMTab working dictionary to a valid OSCAL
    plan-of-action-and-milestones JSON structure.

    Parameters:
        poam          - The internal POAM dictionary (from the form)
        oscal_version - The OSCAL version string to write into the file
                        (e.g. "1.2.2"). Must be supplied explicitly — callers
                        should pass get_oscal_version() from the toolbar callback.
        save_path     - The filesystem path where this POA&M will be saved.
                        Used to compute a portable relative URI for import-ssp.href
                        instead of storing an absolute OS path.
    """
    if oscal_version is None:
        raise ValueError(
            "oscal_version must be provided explicitly (e.g. '1.2.2'). "
            "Pass get_oscal_version() from the toolbar callback."
        )

    # _UUID_RE is compiled once at module level — no need to re-compile here.
    # See the module-level constant defined near the top of this file. (L9 fix)
    doc = {
        "plan-of-action-and-milestones": {
            "uuid": poam.get("uuid") or new_uuid(),
            "metadata": {
                "title":         poam.get("title", ""),
                "last-modified": now_iso(),
                "version":       poam.get("version", "1.0"),
                "oscal-version": oscal_version,
            },
        }
    }
    root = doc["plan-of-action-and-milestones"]

    if poam.get("import_ssp"):
        ssp_raw = poam["import_ssp"].strip()
        # Convert an absolute OS path to a portable relative URI-reference.
        # If the caller provides the save_path we compute a relative path from
        # the POA&M file's directory to the SSP file.  If both files are on
        # different drives (Windows) or the path cannot be made relative we
        # fall back to a file:// URI so the href is at least syntactically valid.
        href = ssp_raw
        if ssp_raw and save_path:
            try:
                ssp_p  = Path(ssp_raw).resolve()
                poam_p = Path(save_path).resolve().parent
                href   = ssp_p.relative_to(poam_p).as_posix()
            except ValueError:
                # Different drive or already relative — keep as-is
                p = Path(ssp_raw)
                href = p.as_uri() if p.is_absolute() else ssp_raw
        root["import-ssp"] = {"href": href}

    if poam.get("system_id"):
        sid = poam["system_id"].strip()
        # Use the RFC 4122 identifier-type only when the value is actually a UUID.
        # For plain system names use a generic URL so the schema enum is satisfied.
        if _UUID_RE.match(sid):
            id_type = "https://ietf.org/rfc/rfc4122"
        else:
            id_type = "https://oscal-user-toolkit/system-id"
        root["system-id"] = {"identifier-type": id_type, "id": sid}

    # observations
    if poam.get("observations"):
        root["observations"] = []
        for o in poam["observations"]:
            # Filter out empty method strings from the list
            obs_methods = [m for m in o.get("methods", []) if m]
            # OSCAL schema requires at least one method. Fall back to "UNKNOWN"
            # if the user cleared all methods, rather than writing an invalid
            # empty array. (M6 fix)
            entry = {
                "uuid":        o.get("uuid") or new_uuid(),
                "description": o.get("description", ""),
                "methods":     obs_methods if obs_methods else ["UNKNOWN"],
                "collected":   o.get("collected", now_iso()),
            }
            if o.get("title"):
                entry["title"] = o["title"]
            if o.get("types"):
                entry["types"] = [t for t in o["types"] if t]
            if o.get("expires"):
                entry["expires"] = o["expires"]
            if o.get("assessed_by"):
                entry["props"] = [{"name": "assessed-by",
                                   "value": o["assessed_by"],
                                   "ns": "https://oscal-user-toolkit/ns/poam"}]
            if o.get("relevant_evidence"):
                entry["relevant-evidence"] = [
                    {"href": e.get("href", ""), "description": e.get("description", "")}
                    for e in o["relevant_evidence"] if e.get("href") or e.get("description")
                ]
            if o.get("remarks"):
                entry["remarks"] = o["remarks"]
            root["observations"].append(entry)

    _CIA_SYSTEM = "https://oscal-user-toolkit/ns/cia-impact"
    # risks
    if poam.get("risks"):
        root["risks"] = []
        for r in poam["risks"]:
            entry = {
                "uuid":        r.get("uuid") or new_uuid(),
                "title":       r.get("title", ""),
                "description": r.get("description", ""),
                "statement":   r.get("statement", ""),
                "status":      r.get("status", "open"),
            }
            if r.get("deadline"):
                entry["deadline"] = r["deadline"]
            cia = {k: r.get(k, "") for k in ("cia_c", "cia_i", "cia_a")}
            if any(cia.values()):
                # Build the CIA facets — Confidentiality, Integrity, Availability
                cia_facets = [
                    {"name": dim, "system": _CIA_SYSTEM, "value": val}
                    for dim, val in [
                        ("confidentiality", cia["cia_c"]),
                        ("integrity",       cia["cia_i"]),
                        ("availability",    cia["cia_a"]),
                    ] if val
                ]
                entry["characterizations"] = [{
                    # OSCAL 1.2.2 requires an 'origin' on every characterization.
                    # We declare the toolkit itself (identified by the POA&M UUID)
                    # as the assessing tool. (H2 fix)
                    "origin": {
                        "actors": [{
                            "type":       "tool",
                            "actor-uuid": root.get("uuid", entry["uuid"]),
                        }]
                    },
                    "facets": cia_facets,
                }]
            if r.get("remediations"):
                entry["remediations"] = []
                for rem in r["remediations"]:
                    rem_entry = {
                        "uuid":        rem.get("uuid") or new_uuid(),
                        "lifecycle":   rem.get("lifecycle", "recommendation"),
                        "title":       rem.get("title", ""),
                        "description": rem.get("description", ""),
                    }
                    if rem.get("remarks"):
                        rem_entry["remarks"] = rem["remarks"]
                    entry["remediations"].append(rem_entry)
            if r.get("remarks"):
                entry["remarks"] = r["remarks"]
            root["risks"].append(entry)

    # findings
    if poam.get("findings"):
        root["findings"] = []
        for f in poam["findings"]:
            entry = {
                "uuid":        f.get("uuid") or new_uuid(),
                "title":       f.get("title", ""),
                "description": f.get("description", ""),
                "target": {
                    "type":      f.get("target_type", "statement-id"),
                    "target-id": f.get("target_id", ""),
                    "status": {
                        "state": f.get("status_state", "not-satisfied"),
                    },
                },
            }
            if f.get("status_reason"):
                entry["target"]["status"]["reason"] = f["status_reason"]
            if f.get("remarks"):
                entry["remarks"] = f["remarks"]
            root["findings"].append(entry)

    # poam-items (required)
    _POAM_NS = "https://oscal-user-toolkit/ns/poam"
    items = poam.get("poam_items", [])
    root["poam-items"] = []
    for item in items:
        pi = {
            "uuid":        item.get("uuid") or new_uuid(),
            "title":       item.get("title", ""),
            "description": item.get("description", ""),
        }
        if item.get("scheduled_completion"):
            pi["props"] = [{"name": "scheduled-completion-date",
                            "value": item["scheduled_completion"],
                            "ns": _POAM_NS}]
        if item.get("related_observation_uuids"):
            pi["related-observations"] = [
                {"observation-uuid": u}
                for u in item["related_observation_uuids"] if u
            ]
        if item.get("related_risk_uuids"):
            pi["related-risks"] = [
                {"risk-uuid": u}
                for u in item["related_risk_uuids"] if u
            ]
        if item.get("related_finding_uuids"):
            pi["related-findings"] = [
                {"finding-uuid": u}
                for u in item["related_finding_uuids"] if u
            ]
        if item.get("remarks"):
            pi["remarks"] = item["remarks"]
        root["poam-items"].append(pi)

    return doc


# =============================================================================
# SHARED PARSE/BUILD HELPERS  (L5, L6, L7 fixes)
# These module-level helpers are used by both POA&M and AR functions so
# the same internal dict structure is always produced for observations and
# risks, regardless of which document type is being read or written.
# =============================================================================

def _prop_value(props_list, name):
    """
    Return the value of the first prop with the given name, or ''.

    OSCAL objects can carry arbitrary metadata as 'props' — a list of
    {"name": "...", "value": "..."} dicts. This helper searches that list.

    Parameters:
        props_list - A list of prop dicts (may be empty)
        name       - The prop name to look for, e.g. "assessed-by"

    Returns:
        The matching value string, or '' if not found.
    """
    return next(
        (p.get("value", "") for p in props_list if p.get("name") == name),
        "",
    )


def _parse_oscal_observation(o):
    """
    Convert a raw OSCAL observation JSON dict into the internal UI model format.

    This helper is shared between parse_poam_file() and parse_ar_file()
    so both document types produce the same internal dict structure. (L7 fix)

    Parameters:
        o - a single observation dict from the raw OSCAL JSON

    Returns:
        An internal observation dict compatible with the POA&M and AR UIs.
    """
    # Read the 'assessed-by' value from the observation's props array, if
    # present. Props are key-value metadata attached to OSCAL objects.
    assessed_by = _prop_value(o.get("props", []), "assessed-by")
    return {
        "uuid":        o.get("uuid", new_uuid()),
        "description": o.get("description", ""),
        "methods":     list(o.get("methods", [])),
        "type":        (o.get("types") or [""])[0],
        "assessed_by": assessed_by,
    }


def parse_poam_file(data):
    """
    Parse a raw OSCAL POA&M JSON dict (as loaded from disk) into the internal
    working dictionary format used by POAMTab.
    """
    root = data.get("plan-of-action-and-milestones", {})
    meta = root.get("metadata", {})

    poam = empty_poam()
    poam["uuid"]    = root.get("uuid", new_uuid())
    poam["title"]   = meta.get("title", "")
    poam["version"] = meta.get("version", "1.0")

    issp = root.get("import-ssp", {})
    poam["import_ssp"] = issp.get("href", "") if isinstance(issp, dict) else ""

    sid = root.get("system-id", {})
    poam["system_id"] = sid.get("id", "") if isinstance(sid, dict) else ""

    # observations — use the module-level _prop_value helper (L7 fix)
    for o in root.get("observations", []):
        ev = [
            {"href": e.get("href", ""), "description": e.get("description", "")}
            for e in o.get("relevant-evidence", [])
        ]
        # Start with the shared minimal fields from _parse_oscal_observation,
        # then merge in the POA&M-specific extra fields (title, types, dates,
        # relevant-evidence, remarks) that the POA&M editor also uses.
        base = _parse_oscal_observation(o)
        poam["observations"].append({
            **base,
            "title":             o.get("title", ""),
            "types":             list(o.get("types", [])),
            "collected":         o.get("collected", ""),
            "expires":           o.get("expires", ""),
            "relevant_evidence": ev,
            "remarks":           o.get("remarks", ""),
        })

    _CIA_SYSTEM = "https://oscal-user-toolkit/ns/cia-impact"
    # risks
    for r in root.get("risks", []):
        rems = []
        for rem in r.get("remediations", []):
            rems.append({
                "uuid":        rem.get("uuid", new_uuid()),
                "lifecycle":   rem.get("lifecycle", "recommendation"),
                "title":       rem.get("title", ""),
                "description": rem.get("description", ""),
                "remarks":     rem.get("remarks", ""),
            })
        cia_c = cia_i = cia_a = ""
        for char in r.get("characterizations", []):
            for facet in char.get("facets", []):
                if facet.get("system") == _CIA_SYSTEM:
                    n, v = facet.get("name", ""), facet.get("value", "")
                    if n == "confidentiality": cia_c = v
                    elif n == "integrity":     cia_i = v
                    elif n == "availability":  cia_a = v
        poam["risks"].append({
            "uuid":        r.get("uuid", new_uuid()),
            "title":       r.get("title", ""),
            "description": r.get("description", ""),
            "statement":   r.get("statement", ""),
            "status":      r.get("status", "open"),
            "deadline":    r.get("deadline", ""),
            "cia_c":       cia_c,
            "cia_i":       cia_i,
            "cia_a":       cia_a,
            "remediations": rems,
            "remarks":     r.get("remarks", ""),
        })

    # findings
    for f in root.get("findings", []):
        tgt    = f.get("target", {})
        status = tgt.get("status", {})
        poam["findings"].append({
            "uuid":         f.get("uuid", new_uuid()),
            "title":        f.get("title", ""),
            "description":  f.get("description", ""),
            "target_type":  tgt.get("type", "statement-id"),
            "target_id":    tgt.get("target-id", ""),
            "status_state": status.get("state", "not-satisfied"),
            "status_reason": status.get("reason", ""),
            "remarks":      f.get("remarks", ""),
        })

    # poam-items
    for item in root.get("poam-items", []):
        poam["poam_items"].append({
            "uuid":        item.get("uuid", new_uuid()),
            "title":       item.get("title", ""),
            "description": item.get("description", ""),
            "scheduled_completion": _prop_value(item, "scheduled-completion-date"),
            "related_observation_uuids": [
                ro.get("observation-uuid", "")
                for ro in item.get("related-observations", [])
            ],
            "related_risk_uuids": [
                rr.get("risk-uuid", "")
                for rr in item.get("related-risks", [])
            ],
            "related_finding_uuids": [
                rf.get("finding-uuid", "")
                for rf in item.get("related-findings", [])
            ],
            "remarks": item.get("remarks", ""),
        })

    return poam


# =============================================================================
# ASSESSMENT PLAN DATA FUNCTIONS
# =============================================================================

def empty_ap():
    """Return a blank Assessment Plan working dictionary used by APTab."""
    return {
        "uuid":                    new_uuid(),
        "title":                   "",
        "version":                 "1.0",
        "import_ssp":              "",   # href to the SSP being assessed
        "reviewed_controls_all":   True, # True → include-all; False → specific IDs
        "reviewed_control_ids":    [],   # list of control IDs (when not include-all)
        "tasks":                   [],   # list of task dicts
    }


def build_oscal_ap(ap, oscal_version=None, save_path=None):
    """
    Convert the internal APTab working dictionary to a valid OSCAL
    assessment-plan JSON structure.

    Parameters:
        ap            - The internal AP dictionary (from the form)
        oscal_version - The OSCAL version string (e.g. "1.2.2"). Required.
        save_path     - Filesystem path where this AP will be saved, used to
                        compute a portable relative URI for import-ssp.href.
    """
    if oscal_version is None:
        raise ValueError(
            "oscal_version must be provided explicitly (e.g. '1.2.2'). "
            "Pass get_oscal_version() from the toolbar callback."
        )

    doc = {
        "assessment-plan": {
            "uuid": ap.get("uuid") or new_uuid(),
            "metadata": {
                "title":         ap.get("title", ""),
                "last-modified": now_iso(),
                "version":       ap.get("version", "1.0"),
                "oscal-version": oscal_version,
            },
            "import-ssp": {"href": ""},
        }
    }
    root = doc["assessment-plan"]

    # import-ssp — convert absolute path to portable relative URI
    ssp_raw = ap.get("import_ssp", "").strip()
    href = ssp_raw
    if ssp_raw and save_path:
        try:
            ssp_p  = Path(ssp_raw).resolve()
            ap_p   = Path(save_path).resolve().parent
            href   = ssp_p.relative_to(ap_p).as_posix()
        except ValueError:
            p = Path(ssp_raw)
            href = p.as_uri() if p.is_absolute() else ssp_raw
    root["import-ssp"]["href"] = href or ""

    # reviewed-controls
    if ap.get("reviewed_controls_all", True):
        control_selections = [{"include-all": {}}]
    else:
        ids = [cid.strip() for cid in ap.get("reviewed_control_ids", []) if cid.strip()]
        control_selections = [{"include-controls": [{"with-ids": ids}]}] if ids else [{"include-all": {}}]
    root["reviewed-controls"] = {"control-selections": control_selections}

    # OSCAL requires assessment-subjects to define what is being assessed.
    # "include-all" is a shorthand meaning all system components are in scope.
    # Without this field the document fails OSCAL schema validation. (H1 fix)
    root["assessment-subjects"] = [{"type": "component", "include-all": {}}]

    # tasks
    if ap.get("tasks"):
        root["tasks"] = []
        for t in ap["tasks"]:
            task = {
                "uuid":  t.get("uuid") or new_uuid(),
                "type":  t.get("type", "milestone"),
                "title": t.get("title", ""),
            }
            if t.get("description"):
                task["description"] = t["description"]

            ttype = t.get("timing_type", "")
            if ttype == "on-date" and t.get("timing_date"):
                task["timing"] = {"on-date": {"date": t["timing_date"]}}
            elif ttype == "within-date-range" and t.get("timing_start"):
                task["timing"] = {"within-date-range": {
                    "start": t["timing_start"],
                    **({"end": t["timing_end"]} if t.get("timing_end") else {}),
                }}
            elif ttype == "at-frequency" and t.get("timing_period"):
                # OSCAL schema requires period to be an integer and unit to be
                # a string (e.g. "days"). We cast period here and include the
                # unit field that was missing before. Default period = 7 days. (H5)
                task["timing"] = {"at-frequency": {
                    "period": int(t.get("timing_period", 7)),
                    "unit":   t.get("timing_unit", "days"),
                }}

            if t.get("remarks"):
                task["remarks"] = t["remarks"]
            root["tasks"].append(task)

    return doc


def parse_ap_file(data):
    """
    Parse a raw OSCAL assessment-plan JSON dict into the internal working
    dictionary format used by APTab.
    """
    root = data.get("assessment-plan", {})
    meta = root.get("metadata", {})

    ap = empty_ap()
    ap["uuid"]    = root.get("uuid", new_uuid())
    ap["title"]   = meta.get("title", "")
    ap["version"] = meta.get("version", "1.0")
    ap["import_ssp"] = root.get("import-ssp", {}).get("href", "")

    # reviewed-controls
    rc  = root.get("reviewed-controls", {})
    sels = rc.get("control-selections", [{}])
    sel0 = sels[0] if sels else {}
    if "include-all" in sel0:
        ap["reviewed_controls_all"] = True
    else:
        ap["reviewed_controls_all"] = False
        inc = sel0.get("include-controls", [{}])
        ids = inc[0].get("with-ids", []) if inc else []
        ap["reviewed_control_ids"] = list(ids)

    # tasks
    for t in root.get("tasks", []):
        timing     = t.get("timing", {})
        ttype      = ""
        t_date     = ""
        t_start    = ""
        t_end      = ""
        t_period   = ""
        if "on-date" in timing:
            ttype  = "on-date"
            t_date = timing["on-date"].get("date", "")
        elif "within-date-range" in timing:
            ttype    = "within-date-range"
            t_start  = timing["within-date-range"].get("start", "")
            t_end    = timing["within-date-range"].get("end", "")
        elif "at-frequency" in timing:
            ttype    = "at-frequency"
            # Read period back as a string so the Entry widget displays it (H5)
            t_period = str(timing["at-frequency"].get("period", ""))

        # Read timing_unit from the at-frequency object; default to "days" (H5)
        t_unit = timing.get("at-frequency", {}).get("unit", "days")

        ap["tasks"].append({
            "uuid":        t.get("uuid", new_uuid()),
            "type":        t.get("type", "milestone"),
            "title":       t.get("title", ""),
            "description": t.get("description", ""),
            "timing_type":  ttype,
            "timing_date":  t_date,
            "timing_start": t_start,
            "timing_end":   t_end,
            "timing_period": t_period,
            # timing_unit is the OSCAL at-frequency.unit value (e.g. "days") (H5)
            "timing_unit":  t_unit,
            "remarks":     t.get("remarks", ""),
        })

    return ap


# =============================================================================
# ASSESSMENT RESULTS DATA FUNCTIONS
# =============================================================================

def empty_ar():
    """Return a blank Assessment Results working dictionary used by ARTab."""
    return {
        "uuid":               new_uuid(),
        "title":              "",
        "version":            "1.0",
        "import_ap":          "",          # href to the Assessment Plan
        "result_uuid":        new_uuid(),
        "result_title":       "",
        "result_description": "",
        "result_start":       now_iso()[:10],
        "result_end":         "",
        "observations":       [],
        "risks":              [],
        "findings":           [],
        "assessment_log":     [],          # Phase 3 — log entries
    }


def build_oscal_ar(ar, oscal_version=None, save_path=None):
    """
    Convert the internal ARTab working dictionary to a valid OSCAL
    assessment-results JSON structure.

    Parameters:
        ar            - The internal AR dictionary (from the form)
        oscal_version - The OSCAL version string (e.g. "1.2.2"). Required.
        save_path     - Filesystem path where this AR will be saved.
    """
    if oscal_version is None:
        raise ValueError(
            "oscal_version must be provided explicitly (e.g. '1.2.2'). "
            "Pass get_oscal_version() from the toolbar callback."
        )

    doc = {
        "assessment-results": {
            "uuid": ar.get("uuid") or new_uuid(),
            "metadata": {
                "title":         ar.get("title", ""),
                "last-modified": now_iso(),
                "version":       ar.get("version", "1.0"),
                "oscal-version": oscal_version,
            },
            "import-ap": {"href": ""},
            "results":   [],
        }
    }
    root = doc["assessment-results"]

    # import-ap — convert absolute path to portable relative URI
    ap_raw = ar.get("import_ap", "").strip()
    href = ap_raw
    if ap_raw and save_path:
        try:
            ap_p  = Path(ap_raw).resolve()
            ar_p  = Path(save_path).resolve().parent
            href  = ap_p.relative_to(ar_p).as_posix()
        except ValueError:
            p = Path(ap_raw)
            href = p.as_uri() if p.is_absolute() else ap_raw
    root["import-ap"]["href"] = href or ""

    # Build the single result object
    result = {
        "uuid":        ar.get("result_uuid") or new_uuid(),
        "title":       ar.get("result_title", ""),
        "description": ar.get("result_description", ""),
        "start":       ar.get("result_start", now_iso()),
        "reviewed-controls": {"control-selections": [{"include-all": {}}]},
    }
    if ar.get("result_end"):
        result["end"] = ar["result_end"]

    # observations
    if ar.get("observations"):
        result["observations"] = []
        for o in ar["observations"]:
            ar_methods = [m for m in o.get("methods", []) if m]
            # Guard against an empty methods array (same M6 fix as POA&M)
            entry = {
                "uuid":        o.get("uuid") or new_uuid(),
                "description": o.get("description", ""),
                "methods":     ar_methods if ar_methods else ["UNKNOWN"],
                "collected":   o.get("collected", now_iso()),
            }
            if o.get("title"):
                entry["title"] = o["title"]
            if o.get("types"):
                entry["types"] = [t for t in o["types"] if t]
            if o.get("expires"):
                entry["expires"] = o["expires"]
            if o.get("relevant_evidence"):
                entry["relevant-evidence"] = [
                    {"href": e.get("href", ""), "description": e.get("description", "")}
                    for e in o["relevant_evidence"]
                    if e.get("href") or e.get("description")
                ]
            if o.get("remarks"):
                entry["remarks"] = o["remarks"]
            result["observations"].append(entry)

    # risks
    if ar.get("risks"):
        result["risks"] = []
        for r in ar["risks"]:
            entry = {
                "uuid":        r.get("uuid") or new_uuid(),
                "title":       r.get("title", ""),
                "description": r.get("description", ""),
                "statement":   r.get("statement", ""),
                "status":      r.get("status", "open"),
            }
            if r.get("deadline"):
                entry["deadline"] = r["deadline"]
            if r.get("remediations"):
                entry["remediations"] = []
                for rem in r["remediations"]:
                    re_entry = {
                        "uuid":        rem.get("uuid") or new_uuid(),
                        "lifecycle":   rem.get("lifecycle", "recommendation"),
                        "title":       rem.get("title", ""),
                        "description": rem.get("description", ""),
                    }
                    if rem.get("remarks"):
                        re_entry["remarks"] = rem["remarks"]
                    entry["remediations"].append(re_entry)
            if r.get("remarks"):
                entry["remarks"] = r["remarks"]
            result["risks"].append(entry)

    # findings
    if ar.get("findings"):
        result["findings"] = []
        for f in ar["findings"]:
            entry = {
                "uuid":        f.get("uuid") or new_uuid(),
                "title":       f.get("title", ""),
                "description": f.get("description", ""),
                "target": {
                    "type":      f.get("target_type", "statement-id"),
                    "target-id": f.get("target_id", ""),
                    "status": {
                        "state": f.get("status_state", "not-satisfied"),
                    },
                    "implementation-status": {
                        "state": f.get("impl_status", "implemented"),
                    },
                },
            }
            if f.get("status_reason"):
                entry["target"]["status"]["reason"] = f["status_reason"]
            obs_uuids = [u for u in f.get("related_obs_uuids", []) if u]
            if obs_uuids:
                entry["related-observations"] = [
                    {"observation-uuid": u} for u in obs_uuids
                ]
            risk_uuids = [u for u in f.get("related_risk_uuids", []) if u]
            if risk_uuids:
                entry["related-risks"] = [
                    {"risk-uuid": u} for u in risk_uuids
                ]
            if f.get("remarks"):
                entry["remarks"] = f["remarks"]
            result["findings"].append(entry)

    # assessment-log (Phase 3)
    if ar.get("assessment_log"):
        log_entries = []
        for le in ar["assessment_log"]:
            log_entry = {
                "uuid":  le.get("uuid") or new_uuid(),
                "start": le.get("start", now_iso()),
            }
            if le.get("end"):
                log_entry["end"] = le["end"]
            if le.get("description"):
                log_entry["description"] = le["description"]
            if le.get("remarks"):
                log_entry["remarks"] = le["remarks"]
            log_entries.append(log_entry)
        result["assessment-log"] = {"entries": log_entries}

    root["results"].append(result)
    return doc


def parse_ar_file(data):
    """
    Parse a raw OSCAL assessment-results JSON dict into the internal working
    dictionary format used by ARTab.
    """
    root   = data.get("assessment-results", {})
    meta   = root.get("metadata", {})
    result = root.get("results", [{}])[0] if root.get("results") else {}

    ar = empty_ar()
    ar["uuid"]    = root.get("uuid", new_uuid())
    ar["title"]   = meta.get("title", "")
    ar["version"] = meta.get("version", "1.0")
    ar["import_ap"] = root.get("import-ap", {}).get("href", "")

    ar["result_uuid"]        = result.get("uuid", new_uuid())
    ar["result_title"]       = result.get("title", "")
    ar["result_description"] = result.get("description", "")
    ar["result_start"]       = result.get("start", "")[:10] if result.get("start") else ""
    ar["result_end"]         = result.get("end",   "")[:10] if result.get("end")   else ""

    # observations — use the shared _parse_oscal_observation helper (L7 fix)
    for o in result.get("observations", []):
        ev = [
            {"href": e.get("href", ""), "description": e.get("description", "")}
            for e in o.get("relevant-evidence", [])
        ]
        # Get the common fields via the shared helper, then add AR-specific extras
        base = _parse_oscal_observation(o)
        ar["observations"].append({
            **base,
            "title":             o.get("title", ""),
            "types":             list(o.get("types", [])),
            "collected":         o.get("collected", "")[:10] if o.get("collected") else "",
            "expires":           o.get("expires", "")[:10]   if o.get("expires")   else "",
            "relevant_evidence": ev,
            "remarks":           o.get("remarks", ""),
        })

    # risks
    for r in result.get("risks", []):
        rems = []
        for rem in r.get("remediations", []):
            rems.append({
                "uuid":        rem.get("uuid", new_uuid()),
                "lifecycle":   rem.get("lifecycle", "recommendation"),
                "title":       rem.get("title", ""),
                "description": rem.get("description", ""),
                "remarks":     rem.get("remarks", ""),
            })
        ar["risks"].append({
            "uuid":        r.get("uuid", new_uuid()),
            "title":       r.get("title", ""),
            "description": r.get("description", ""),
            "statement":   r.get("statement", ""),
            "status":      r.get("status", "open"),
            "deadline":    r.get("deadline", "")[:10] if r.get("deadline") else "",
            "remediations": rems,
            "remarks":     r.get("remarks", ""),
        })

    # findings
    for f in result.get("findings", []):
        tgt    = f.get("target", {})
        status = tgt.get("status", {})
        impl   = tgt.get("implementation-status", {})
        ar["findings"].append({
            "uuid":          f.get("uuid", new_uuid()),
            "title":         f.get("title", ""),
            "description":   f.get("description", ""),
            "target_type":   tgt.get("type", "statement-id"),
            "target_id":     tgt.get("target-id", ""),
            "status_state":  status.get("state", "not-satisfied"),
            "status_reason": status.get("reason", ""),
            "impl_status":   impl.get("state", "implemented"),
            "related_obs_uuids": [
                ro.get("observation-uuid", "")
                for ro in f.get("related-observations", [])
            ],
            "related_risk_uuids": [
                rr.get("risk-uuid", "")
                for rr in f.get("related-risks", [])
            ],
            "remarks":       f.get("remarks", ""),
        })

    # assessment-log
    log = result.get("assessment-log", {})
    for le in log.get("entries", []):
        ar["assessment_log"].append({
            "uuid":        le.get("uuid", new_uuid()),
            "start":       le.get("start", "")[:10] if le.get("start") else "",
            "end":         le.get("end",   "")[:10] if le.get("end")   else "",
            "description": le.get("description", ""),
            "remarks":     le.get("remarks", ""),
        })

    return ar
